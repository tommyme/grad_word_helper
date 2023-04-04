import docx
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx import document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.styles.style import BaseStyle

font_size = Pt(7.875)  # 10.5磅 = 7.875 Pt


# 打开Word文档
doc: document.Document = docx.Document(
    os.path.join(
        os.environ["HOME"], "Downloads", '毕业论文.docx'
    )
    # 'example_modified.docx'
)


# 更改预设样式
for style in doc.styles:
    pass
    # print(style.name, style.type)


# add_style使用docx.styles.styles.Style的Styles类的add_style方法 调用了StyleFactory创造样式
new_style = doc.styles.add_style("图题", WD_STYLE_TYPE.PARAGRAPH)
new_style.font.size = Pt(10.5)  # 5号字体


def checkPic(para: Paragraph):
    if len(para.runs) == 1:
        if para.runs[0]._element.xpath('.//wp:inline'):
            return True
    return False


def imgAddTitle(doc: document.Document):
    idx = 0
    while idx < len(doc.paragraphs):
        current_para = doc.paragraphs[idx]
        if checkPic(current_para):
            # 插入到图片的下方
            inserted = doc.paragraphs[idx+1].insert_paragraph_before(text="你好hello world", style="图题")
            # 居中
            inserted.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            idx += 2
        else:
            idx += 1

import re
def checkTitleStyle(doc: document.Document):
    for idx, para in enumerate(doc.paragraphs):
        # 匹配三级标题
        if re.match(r"(第\d章|\d\.\d\.\d|\d\.\d).*", para.text):
            if re.findall(r"\t\d+$", para.text):    # 过滤目录页
                continue
            if para.paragraph_format.first_line_indent is not None: # para有缩进
                continue    
            print(para.text)
# doc.save('example_modified.docx')

if __name__ == "__main__":
    checkTitleStyle(doc)
